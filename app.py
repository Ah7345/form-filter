import streamlit as st
import json
import io
import os
from pathlib import Path
from typing import Dict, List, Any

# OpenAI client
from openai import OpenAI

# File processing
import PyPDF2
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT, TA_CENTER

# Arabic text processing
import arabic_reshaper
from bidi.algorithm import get_display

# Font configuration
AR_FONT_REGULAR = "NotoNaskhArabic-Regular"
AR_FONT_BOLD = "NotoNaskhArabic-Bold"
AR_FONT_REGULAR_PATH = "fonts/NotoNaskhArabic-Regular.ttf"
AR_FONT_BOLD_PATH = "fonts/NotoNaskhArabic-Bold.ttf"

# OpenAI API configuration
def get_openai_api_key():
    """Get OpenAI API key from environment or secrets"""
    try:
        # Try to get from Streamlit secrets first
        if hasattr(st, 'secrets') and st.secrets:
            try:
                # Try different ways to access the secret
                api_key = st.secrets.get("OPENAI_API_KEY", "")
                if not api_key:
                    # Try without default section
                    api_key = st.secrets.get("OPENAI_API_KEY", "")
                if not api_key:
                    # Try direct access
                    api_key = getattr(st.secrets, "OPENAI_API_KEY", "")
                
                if api_key and api_key != "your-api-key-here":
                    return api_key
            except Exception:
                pass
        
        # Fallback to environment variable
        return os.getenv("OPENAI_API_KEY", "")
    except Exception:
        return os.getenv("OPENAI_API_KEY", "")

def register_arabic_fonts():
    """Register Arabic fonts for PDF generation"""
    try:
                # Check if font files exist
        if not os.path.exists(AR_FONT_REGULAR_PATH) or not os.path.exists(AR_FONT_BOLD_PATH):
            st.warning("ملفات الخطوط العربية غير موجودة")
            st.info("سيتم استخدام خطوط النظام المتاحة")
            return get_system_fallback_font()
        
        # Try to register the Noto Naskh Arabic fonts
        pdfmetrics.registerFont(TTFont(AR_FONT_REGULAR, AR_FONT_REGULAR_PATH))
        pdfmetrics.registerFont(TTFont(AR_FONT_BOLD, AR_FONT_BOLD_PATH))
        
        st.success("✅ تم تسجيل الخطوط العربية بنجاح!")
        return True
        
    except Exception as e:
        st.warning(f"⚠️ تحذير: لا يمكن تسجيل الخطوط العربية: {str(e)}")
        st.info("💡 سيتم استخدام خطوط النظام المتاحة")
        return get_system_fallback_font()

def get_system_fallback_font():
    """Get the best available system font for Arabic support"""
    system_fonts = [
        # macOS fonts
        ('/System/Library/Fonts/Arial.ttf', 'Arial'),
        ('/System/Library/Fonts/Arial Unicode MS.ttf', 'ArialUnicodeMS'),
        ('/System/Library/Fonts/Helvetica.ttc', 'Helvetica'),
        
        # Linux fonts
        ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 'DejaVuSans'),
        ('/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf', 'LiberationSans'),
        
        # Windows fonts (if running on Windows)
        ('C:/Windows/Fonts/arial.ttf', 'Arial'),
        ('C:/Windows/Fonts/arialuni.ttf', 'ArialUnicodeMS'),
    ]
    
    for font_path, font_name in system_fonts:
        try:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                st.info(f"💡 تم استخدام خط النظام: {font_name}")
                return font_name
        except:
            continue
    
    # Last resort - use default Helvetica
    st.warning("⚠️ استخدام خط النظام: Helvetica")
    return 'Helvetica'

# Page configuration
st.set_page_config(
    page_title="نظام بطاقة الوصف المهني",
    page_icon="",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for RTL and corporate styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@300;400;600;700&display=swap');
    
    .main {
        direction: rtl;
        font-family: 'Cairo', sans-serif;
    }
    
    .stApp {
        direction: rtl;
        font-family: 'Cairo', sans-serif;
    }
    
    .form-header {
        text-align: center;
        color: #1f4e79;
        font-size: 2.5rem;
        font-weight: 700;
        margin-bottom: 2rem;
        padding: 2rem;
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    .section-header {
        background: #e3f2fd;
        color: #1565c0;
        padding: 1rem 1.5rem;
        border-radius: 10px;
        margin: 1.5rem 0 1rem 0;
        border-right: 5px solid #2196f3;
        font-weight: 600;
        font-size: 1.2rem;
    }
    
    .subsection-header {
        background: #f3e5f5;
        color: #7b1fa2;
        padding: 0.75rem 1rem;
        border-radius: 8px;
        margin: 1rem 0 0.5rem 0;
        border-right: 4px solid #9c27b0;
        font-weight: 500;
        font-size: 1.1rem;
    }
    
    .field-group {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        border: 1px solid #e0e0e0;
    }
    
    .repeatable-section {
        background: #fafafa;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        border: 2px dashed #e0e0e0;
    }
    
    .row-controls {
        display: flex;
        gap: 0.5rem;
        margin: 0.5rem 0;
        justify-content: flex-end;
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #007bff 0%, #0056b3 100%);
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 6px;
        font-weight: 500;
        font-size: 0.9rem;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.15);
    }
    
    .add-button {
        background: linear-gradient(135deg, #28a745 0%, #1e7e34 100%) !important;
    }
    
    .remove-button {
        background: linear-gradient(135deg, #dc3545 0%, #c82333 100%) !important;
    }
    
    .submit-button {
        background: linear-gradient(135deg, #6f42c1 0%, #5a2d91 100%) !important;
        padding: 1rem 2rem !important;
        font-size: 1.1rem !important;
    }
    
    .validation-error {
        color: #dc3545;
        font-size: 0.9rem;
        margin-top: 0.25rem;
        font-weight: 500;
    }
    
    .success-message {
        background: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #c3e6cb;
        margin: 1rem 0;
    }
    
    .json-output {
        background: #f8f9fa;
        border: 1px solid #dee2e6;
        border-radius: 8px;
        padding: 1rem;
        font-family: 'Courier New', monospace;
        font-size: 0.9rem;
        direction: ltr;
        text-align: left;
    }
    
    .file-upload-section {
        background: linear-gradient(135deg, #e3f2fd 0%, #f3e5f5 100%);
        padding: 2rem;
        border-radius: 15px;
        margin: 1rem 0;
        border: 2px solid #2196f3;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    
    .ai-analysis-button {
        background: linear-gradient(135deg, #9c27b0 0%, #673ab7 100%) !important;
        color: white !important;
        border: none !important;
        padding: 0.75rem 1.5rem !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        transition: all 0.3s ease !important;
    }
    
    .ai-analysis-button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 12px rgba(0,0,0,0.2) !important;
    }
</style>
""", unsafe_allow_html=True)

def initialize_session_state():
    """Initialize session state for form data"""
    if 'form_data' not in st.session_state:
        st.session_state.form_data = {
            'ref_data': {
                'main_group': '',
                'main_group_code': '',
                'sub_group': '',
                'sub_group_code': '',
                'secondary_group': '',
                'secondary_group_code': '',
                'unit_group': '',
                'unit_group_code': '',
                'job': '',
                'job_code': '',
                'work_location': '',
                'grade': ''
            },
            'summary': '',
            'internal_communications': [{'entity': '', 'purpose': ''}],
            'external_communications': [{'entity': '', 'purpose': ''}],
            'job_levels': [{'level': '', 'code': '', 'role': '', 'progression': ''}],
            'behavioral_competencies': [{'name': '', 'level': ''}],
            'core_competencies': [{'name': '', 'level': ''}],
            'leadership_competencies': [{'name': '', 'level': ''}],
            'technical_competencies': [{'name': '', 'level': ''}],
            'leadership_tasks': [''],
            'specialized_tasks': [''],
            'other_tasks': [''],
            'behavioral_table': [{'number': 1, 'name': '', 'level': ''}],
            'technical_table': [{'number': 1, 'name': '', 'level': ''}],
            'kpis': [{'number': 1, 'metric': '', 'measure': ''}]
        }

def add_row(data_list: List, template: Dict = None):
    """Add a new row to a repeatable section"""
    if template is None:
        template = {}
    data_list.append(template.copy())
    st.rerun()

def remove_row(data_list: List, index: int):
    """Remove a row from a repeatable section"""
    if len(data_list) > 1:
        data_list.pop(index)
        st.rerun()

def extract_text_from_file(uploaded_file):
    """Extract text from uploaded file (PDF, DOCX, or TXT)"""
    try:
        file_extension = Path(uploaded_file.name).suffix.lower()
        
        if file_extension == '.pdf':
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        elif file_extension == '.docx':
            doc = docx.Document(uploaded_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        elif file_extension == '.txt':
            return str(uploaded_file.read(), "utf-8")
        
        else:
            return None
            
    except Exception as e:
        st.error(f"خطأ في قراءة الملف: {str(e)}")
        return None

def analyze_job_description_with_ai(text_content):
    """Use OpenAI to analyze job description and extract relevant information"""
    # Check if API key is available
    api_key = get_openai_api_key()
    if not api_key or api_key == "your-api-key-here":
        st.error("❌ مفتاح API الخاص بـ OpenAI غير متوفر")
        st.info("💡 يرجى إضافة مفتاح API في متغيرات البيئة أو ملف Streamlit secrets")
        return None
    
    try:
        # Show progress
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        status_text.text("🔄 جاري إعداد طلب AI...")
        progress_bar.progress(20)
        
        system_prompt = """You are an expert in analyzing job descriptions. Analyze the provided text and extract the following information in a structured JSON format.

CRITICAL: You must return ONLY a valid JSON object with this exact structure. Do not include any explanations, markdown formatting, or additional text.

{
  "ref_data": {
    "main_group": "string or empty string",
    "main_group_code": "string or empty string", 
    "sub_group": "string or empty string",
    "sub_group_code": "string or empty string",
    "secondary_group": "string or empty string",
    "secondary_group_code": "string or empty string",
    "unit_group": "string or empty string",
    "unit_group_code": "string or empty string",
    "job": "string or empty string",
    "job_code": "string or empty string",
    "work_location": "string or empty string",
    "grade": "string or empty string"
  },
  "summary": "string or empty string",
  "internal_communications": [
    {"entity": "string or empty string", "purpose": "string or empty string"}
  ],
  "external_communications": [
    {"entity": "string or empty string", "purpose": "string or empty string"}
  ],
  "job_levels": [
    {"level": "string or empty string", "code": "string or empty string", "role": "string or empty string", "progression": "string or empty string"}
  ],
  "behavioral_competencies": [
    {"name": "string or empty string", "level": "string or empty string"}
  ],
  "core_competencies": [
    {"name": "string or empty string", "level": "string or empty string"}
  ],
  "leadership_competencies": [
    {"name": "string or empty string", "level": "string or empty string"}
  ],
  "technical_competencies": [
    {"name": "string or empty string", "level": "string or empty string"}
  ],
  "leadership_tasks": ["string or empty string"],
  "specialized_tasks": ["string or empty string"],
  "other_tasks": ["string or empty string"],
  "behavioral_table": [
    {"number": 1, "name": "string or empty string", "level": "string or empty string"}
  ],
  "technical_table": [
    {"number": 1, "name": "string or empty string", "level": "string or empty string"}
  ],
  "kpis": [
    {"number": 1, "metric": "string or empty string", "measure": "string or empty string"}
  ]
}

RULES:
1. Return ONLY the JSON object
2. Use empty strings for missing information
3. Ensure all arrays have at least one item
4. Use Arabic text for values when appropriate
5. No markdown, no explanations, no additional text"""
        
        user_prompt = f"Analyze this job description text and extract the information in the exact JSON format specified:\n\n{text_content}"
        
        status_text.text("🤖 جاري إرسال الطلب إلى OpenAI...")
        progress_bar.progress(40)
        
        client = OpenAI(api_key=get_openai_api_key())
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=3000,
            temperature=0.1
        )
        
        status_text.text("✅ تم استلام الرد من AI...")
        progress_bar.progress(80)
        
        result = response.choices[0].message.content.strip()
        
        # Try to clean the response to extract JSON
        if result.startswith('```json'):
            result = result.replace('```json', '').replace('```', '').strip()
        elif result.startswith('```'):
            result = result.replace('```', '').strip()
        
        status_text.text("✅ تم الانتهاء من التحليل!")
        progress_bar.progress(100)
        
        # Clear progress indicators
        progress_bar.empty()
        status_text.empty()
        
        return result
        
    except Exception as e:
        error_msg = str(e)
        if "authentication" in error_msg.lower() or "401" in error_msg:
                            st.error("خطأ في مصادقة OpenAI API. تأكد من صحة مفتاح API.")
        elif "rate limit" in error_msg.lower() or "429" in error_msg:
                            st.error("تم تجاوز حد الطلبات. يرجى الانتظار قليلاً والمحاولة مرة أخرى.")
        elif "api" in error_msg.lower():
                            st.error(f"خطأ في API: {error_msg}")
        else:
                            st.error(f"خطأ غير متوقع: {error_msg}")
        return None

def auto_fill_form_with_ai(ai_analysis):
    """Auto-fill the form with AI analysis results"""
    try:
        # Try to parse the AI response as JSON
        if ai_analysis and ai_analysis.strip().startswith('{'):
            try:
                parsed_data = json.loads(ai_analysis)
                
                # Show what was extracted
                st.success("تم تحليل النص بنجاح! جاري ملء النموذج...")
                
                # Show raw AI response first
                with st.expander("الرد الخام من AI", expanded=False):
                    st.code(ai_analysis, language="json")
                
                # Display extracted information in a nice format
                with st.expander("المعلومات المستخرجة من AI", expanded=True):
                    st.markdown("---")
                    
                    # Show structured information
                    if 'ref_data' in parsed_data:
                        st.subheader("البيانات المرجعية")
                        ref_data = parsed_data['ref_data']
                        cols = st.columns(2)
                        for i, (key, value) in enumerate(ref_data.items()):
                            if value:  # Only show non-empty values
                                with cols[i % 2]:
                                    st.metric(label=key, value=value)
                    
                    if 'summary' in parsed_data and parsed_data['summary']:
                        st.subheader("ملخص الوظيفة")
                        st.info(parsed_data['summary'])
                    
                    if 'internal_communications' in parsed_data:
                        st.subheader("قنوات التواصل الداخلية")
                        for comm in parsed_data['internal_communications']:
                            if comm.get('entity') or comm.get('purpose'):
                                st.write(f"• **{comm.get('entity', '')}** - {comm.get('purpose', '')}")
                    
                    if 'external_communications' in parsed_data:
                        st.subheader("قنوات التواصل الخارجية")
                        for comm in parsed_data['external_communications']:
                            if comm.get('entity') or comm.get('purpose'):
                                st.write(f"• **{comm.get('entity', '')}** - {comm.get('purpose', '')}")
                    
                    if 'job_levels' in parsed_data:
                        st.subheader("مستويات الوظيفة")
                        for level in parsed_data['job_levels']:
                            if any(level.values()):
                                st.write(f"• **{level.get('level', '')}** - {level.get('role', '')} - {level.get('progression', '')}")
                    
                    if 'behavioral_competencies' in parsed_data:
                        st.subheader("الكفاءات السلوكية")
                        for comp in parsed_data['behavioral_competencies']:
                            if any(comp.values()):
                                st.write(f"• **{comp.get('name', '')}** - المستوى: {comp.get('level', '')}")
                    
                    if 'core_competencies' in parsed_data:
                        st.subheader("الكفاءات الأساسية")
                        for comp in parsed_data['core_competencies']:
                            if any(comp.values()):
                                st.write(f"• **{comp.get('name', '')}** - المستوى: {comp.get('level', '')}")
                    
                    if 'leadership_competencies' in parsed_data:
                        st.subheader("الكفاءات القيادية")
                        for comp in parsed_data['leadership_competencies']:
                            if any(comp.values()):
                                st.write(f"• **{comp.get('name', '')}** - المستوى: {comp.get('level', '')}")
                    
                    if 'technical_competencies' in parsed_data:
                        st.subheader("الكفاءات التقنية")
                        for comp in parsed_data['technical_competencies']:
                            if any(comp.values()):
                                st.write(f"• **{comp.get('name', '')}** - المستوى: {comp.get('level', '')}")
                    
                    if 'leadership_tasks' in parsed_data:
                        st.subheader("المهام القيادية")
                        for task in parsed_data['leadership_tasks']:
                            if task:
                                st.write(f"• {task}")
                    
                    if 'specialized_tasks' in parsed_data:
                        st.subheader("المهام المتخصصة")
                        for task in parsed_data['specialized_tasks']:
                            if task:
                                st.write(f"• {task}")
                    
                    if 'other_tasks' in parsed_data:
                        st.subheader("المهام الأخرى")
                        for task in parsed_data['other_tasks']:
                            if task:
                                st.write(f"• {task}")
                    
                    if 'kpis' in parsed_data:
                        st.subheader("مؤشرات الأداء الرئيسية")
                        for kpi in parsed_data['kpis']:
                            if any(kpi.values()):
                                st.write(f"• **{kpi.get('metric', '')}** - {kpi.get('measure', '')}")
                
                # Update form data with AI results
                if 'ref_data' in parsed_data:
                    st.session_state.form_data['ref_data'].update(parsed_data['ref_data'])
                if 'summary' in parsed_data:
                    st.session_state.form_data['summary'] = parsed_data['summary']
                if 'internal_communications' in parsed_data:
                    st.session_state.form_data['internal_communications'] = parsed_data['internal_communications']
                if 'external_communications' in parsed_data:
                    st.session_state.form_data['external_communications'] = parsed_data['external_communications']
                if 'job_levels' in parsed_data:
                    st.session_state.form_data['job_levels'] = parsed_data['job_levels']
                if 'behavioral_competencies' in parsed_data:
                    st.session_state.form_data['behavioral_competencies'] = parsed_data['behavioral_competencies']
                if 'core_competencies' in parsed_data:
                    st.session_state.form_data['core_competencies'] = parsed_data['core_competencies']
                if 'leadership_competencies' in parsed_data:
                    st.session_state.form_data['leadership_competencies'] = parsed_data['leadership_competencies']
                if 'technical_competencies' in parsed_data:
                    st.session_state.form_data['technical_competencies'] = parsed_data['technical_competencies']
                if 'leadership_tasks' in parsed_data:
                    st.session_state.form_data['leadership_tasks'] = parsed_data['leadership_tasks']
                if 'specialized_tasks' in parsed_data:
                    st.session_state.form_data['specialized_tasks'] = parsed_data['specialized_tasks']
                if 'other_tasks' in parsed_data:
                    st.session_state.form_data['other_tasks'] = parsed_data['other_tasks']
                if 'behavioral_table' in parsed_data:
                    st.session_state.form_data['behavioral_table'] = parsed_data['behavioral_table']
                if 'technical_table' in parsed_data:
                    st.session_state.form_data['technical_table'] = parsed_data['technical_table']
                if 'kpis' in parsed_data:
                    st.session_state.form_data['kpis'] = parsed_data['kpis']
                
                st.success("تم ملء النموذج تلقائياً باستخدام تحليل AI!")
                st.info("يمكنك الآن مراجعة وتعديل البيانات حسب الحاجة")
                
                # Store AI analysis for PDF generation
                st.session_state['last_ai_analysis'] = ai_analysis
                
                # Show summary of what was filled
                st.markdown("### ملخص ما تم ملؤه:")
                summary_items = []
                if 'ref_data' in parsed_data:
                    filled_refs = sum(1 for v in parsed_data['ref_data'].values() if v)
                    summary_items.append(f"• {filled_refs} من البيانات المرجعية")
                if 'summary' in parsed_data and parsed_data['summary']:
                    summary_items.append("• ملخص الوظيفة")
                if 'internal_communications' in parsed_data:
                    filled_comms = sum(1 for c in parsed_data['internal_communications'] if any(c.values()))
                    if filled_comms > 0:
                        summary_items.append(f"• {filled_comms} قناة تواصل داخلية")
                if 'external_communications' in parsed_data:
                    filled_comms = sum(1 for c in parsed_data['external_communications'] if any(c.values()))
                    if filled_comms > 0:
                        summary_items.append(f"• {filled_comms} قناة تواصل خارجية")
                if 'job_levels' in parsed_data:
                    filled_levels = sum(1 for l in parsed_data['job_levels'] if any(l.values()))
                    if filled_levels > 0:
                        summary_items.append(f"• {filled_levels} مستوى وظيفي")
                if 'behavioral_competencies' in parsed_data:
                    filled_comps = sum(1 for c in parsed_data['behavioral_competencies'] if any(c.values()))
                    if filled_comps > 0:
                        summary_items.append(f"• {filled_comps} كفاءة سلوكية")
                if 'core_competencies' in parsed_data:
                    filled_comps = sum(1 for c in parsed_data['core_competencies'] if any(c.values()))
                    if filled_comps > 0:
                        summary_items.append(f"• {filled_comps} كفاءة أساسية")
                if 'leadership_competencies' in parsed_data:
                    filled_comps = sum(1 for c in parsed_data['leadership_competencies'] if any(c.values()))
                    if filled_comps > 0:
                        summary_items.append(f"• {filled_comps} كفاءة قيادية")
                if 'technical_competencies' in parsed_data:
                    filled_comps = sum(1 for c in parsed_data['technical_competencies'] if any(c.values()))
                    if filled_comps > 0:
                        summary_items.append(f"• {filled_comps} كفاءة تقنية")
                if 'leadership_tasks' in parsed_data:
                    filled_tasks = sum(1 for t in parsed_data['leadership_tasks'] if t)
                    if filled_tasks > 0:
                        summary_items.append(f"• {filled_tasks} مهمة قيادية")
                if 'specialized_tasks' in parsed_data:
                    filled_tasks = sum(1 for t in parsed_data['specialized_tasks'] if t)
                    if filled_tasks > 0:
                        summary_items.append(f"• {filled_tasks} مهمة متخصصة")
                if 'other_tasks' in parsed_data:
                    filled_tasks = sum(1 for t in parsed_data['other_tasks'] if t)
                    if filled_tasks > 0:
                        summary_items.append(f"• {filled_tasks} مهمة أخرى")
                if 'kpis' in parsed_data:
                    filled_kpis = sum(1 for k in parsed_data['kpis'] if any(k.values()))
                    if filled_kpis > 0:
                        summary_items.append(f"• {filled_kpis} مؤشر أداء")
                
                for item in summary_items:
                    st.write(item)
                
                # Offer to save AI analysis
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.button("💾 حفظ تحليل AI", key="save_ai_analysis"):
                        try:
                            # Create filename with timestamp
                            from datetime import datetime
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            filename = f"تحليل_AI_{timestamp}.json"
                            
                            # Save AI analysis
                            st.download_button(
                                label="📥 تحميل تحليل AI",
                                data=ai_analysis,
                                file_name=filename,
                                mime="application/json"
                            )
                            st.success(f"✅ تم حفظ تحليل AI في ملف: {filename}")
                        except Exception as e:
                            st.error(f"❌ خطأ في حفظ الملف: {str(e)}")
                
                with col2:
                    if st.button("إنشاء تقرير PDF", key="ai_pdf_report"):
                        try:
                            with st.spinner("جاري إنشاء تقرير PDF..."):
                                # Generate PDF with AI analysis
                                pdf_content = generate_pdf_report(st.session_state.form_data, ai_analysis)
                                
                                if pdf_content:
                                    # Create filename with timestamp
                                    from datetime import datetime
                                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                    filename = f"تقرير_AI_{timestamp}.pdf"
                                    
                                    # Download button
                                    st.download_button(
                                        label="📥 تحميل التقرير PDF",
                                        data=pdf_content,
                                        file_name=filename,
                                        mime="application/pdf"
                                    )
                                    st.success(f"تم إنشاء التقرير بنجاح!")
                                else:
                                    st.error("فشل في إنشاء التقرير PDF")
                        except Exception as e:
                            st.error(f"خطأ في إنشاء التقرير: {str(e)}")
                
                with col3:
                    if st.button("إنشاء تقرير DOCX", key="ai_docx_report"):
                        try:
                            with st.spinner("جاري إنشاء تقرير DOCX..."):
                                # Generate DOCX with AI analysis
                                docx_content = generate_docx_report(st.session_state.form_data, ai_analysis)
                                
                                if docx_content:
                                    # Create filename with timestamp
                                    from datetime import datetime
                                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                    filename = f"تقرير_AI_{timestamp}.docx"
                                    
                                    # Download button
                                    st.download_button(
                                        label="تحميل التقرير DOCX",
                                        data=docx_content,
                                        file_name=filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                                    st.success(f"تم إنشاء التقرير DOCX بنجاح!")
                                else:
                                    st.error("فشل في إنشاء التقرير DOCX")
                        except Exception as e:
                            st.error(f"خطأ في إنشاء التقرير: {str(e)}")
                
                st.rerun()
                
            except json.JSONDecodeError as e:
                st.error(f"خطأ في تحليل JSON: {str(e)}")
                st.warning("لم يتمكن AI من إرجاع بيانات منظمة. سيتم عرض التحليل النصي.")
                st.text_area("تحليل AI:", value=ai_analysis, height=200)
                st.info("حاول مرة أخرى أو استخدم نصاً أوضح")
                
                # Debug: Show what AI actually returned
                st.markdown("### تصحيح الأخطاء:")
                st.code(f"AI Response: {ai_analysis[:500]}...", language="text")
                
                # Offer retry with simplified prompt
                if st.button("إعادة المحاولة مع تلميح مبسط", key="retry_simple"):
                    st.info("جاري إعادة المحاولة مع تلميح مبسط...")
                    simple_prompt = """Return ONLY a valid JSON object with this structure:
{
  "ref_data": {"main_group": "", "job": "", "work_location": ""},
  "summary": "",
  "internal_communications": [{"entity": "", "purpose": ""}],
  "external_communications": [{"entity": "", "purpose": ""}],
  "job_levels": [{"level": "", "role": ""}],
  "behavioral_competencies": [{"name": "", "level": ""}],
  "core_competencies": [{"name": "", "level": ""}],
  "leadership_competencies": [{"name": "", "level": ""}],
  "technical_competencies": [{"name": "", "level": ""}],
  "leadership_tasks": [""],
  "specialized_tasks": [""],
  "other_tasks": [""],
  "behavioral_table": [{"number": 1, "name": "", "level": ""}],
  "technical_table": [{"number": 1, "name": "", "level": ""}],
  "kpis": [{"number": 1, "metric": "", "measure": ""}]
}"""
                    
                    try:
                        # Get the original text from session state
                        original_text = st.session_state.get('last_analyzed_text', '')
                        if not original_text:
                            st.error("لا يمكن العثور على النص الأصلي")
                            return
                        
                        client = openai.OpenAI(api_key=get_openai_api_key())
                        retry_response = client.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": simple_prompt},
                                {"role": "user", "content": f"Analyze: {original_text}"}
                            ],
                            max_tokens=2000,
                            temperature=0.1
                        )
                        
                        retry_result = retry_response.choices[0].message.content.strip()
                        if retry_result.startswith('```'):
                            retry_result = retry_result.replace('```json', '').replace('```', '').strip()
                        
                        st.success("تم إعادة المحاولة!")
                        auto_fill_form_with_ai(retry_result)
                        
                    except Exception as retry_e:
                        st.error(f"فشل في إعادة المحاولة: {str(retry_e)}")
        else:
            st.warning("لم يتمكن AI من إرجاع بيانات منظمة. سيتم عرض التحليل النصي.")
            st.text_area("تحليل AI:", value=ai_analysis, height=200)
            st.info("حاول مرة أخرى أو استخدم نصاً أوضح")
            
    except Exception as e:
        st.error(f"خطأ في ملء النموذج: {str(e)}")
        st.info("يرجى المحاولة مرة أخرى")

def process_arabic_text(text):
    """Process Arabic text for proper display in PDF"""
    if not text or not isinstance(text, str):
        return text
    
    try:
        # Reshape Arabic text
        reshaped_text = arabic_reshaper.reshape(text)
        # Apply bidirectional algorithm for RTL text
        bidi_text = get_display(reshaped_text)
        return bidi_text
    except:
        return text

def A(text):
    """Short alias for process_arabic_text to keep code tidy"""
    return process_arabic_text(text)

def generate_docx_report(form_data, ai_analysis=None):
    """Generate a professional DOCX form template from form data"""
    try:
        # Create a new Word document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "نموذج بطاقة الوصف المهني"
        doc.core_properties.author = "نظام بطاقة الوصف المهني"
        
        # Add main title
        title = doc.add_heading("نموذج بطاقة الوصف المهني", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        timestamp_para = doc.add_paragraph(f"تاريخ الإنشاء: {timestamp}")
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add spacing
        
        # ===== PART A: نموذج بطاقة الوصف المهني =====
        doc.add_heading("أ- نموذج بطاقة الوصف المهني", level=1)
        doc.add_paragraph()
        
        # 1. Reference Data Section
        doc.add_heading("1. البيانات المرجعية للمهنة", level=2)
        ref_data = form_data.get('ref_data', {})
        
        # Create reference data table
        ref_table = doc.add_table(rows=1, cols=2)
        ref_table.style = 'Table Grid'
        ref_table.rows[0].cells[0].text = "المجال"
        ref_table.rows[0].cells[1].text = "القيمة"
        
        # Add reference data rows
        ref_fields = [
            ('المجموعة الرئيسية', ref_data.get('main_group', '')),
            ('رمز المجموعة الرئيسية', ref_data.get('main_group_code', '')),
            ('المجموعة الفرعية', ref_data.get('sub_group', '')),
            ('رمز المجموعة الفرعية', ref_data.get('sub_group_code', '')),
            ('المجموعة الثانوية', ref_data.get('secondary_group', '')),
            ('رمز المجموعة الثانوية', ref_data.get('secondary_group_code', '')),
            ('مجموعة الوحدة', ref_data.get('unit_group', '')),
            ('رمز مجموعة الوحدة', ref_data.get('unit_group_code', '')),
            ('المهنة', ref_data.get('job', '')),
            ('رمز المهنة', ref_data.get('job_code', '')),
            ('موقع العمل', ref_data.get('work_location', '')),
            ('الدرجة', ref_data.get('grade', ''))
        ]
        
        for field_name, field_value in ref_fields:
            row = ref_table.add_row()
            row.cells[0].text = field_name
            row.cells[1].text = field_value if field_value else "_________________"
        
        doc.add_paragraph()  # Add spacing
        
        # 2. General Summary Section
        doc.add_heading("2. الملخص العام للمهنة", level=2)
        summary = form_data.get('summary', '')
        if summary:
            doc.add_paragraph(summary)
        else:
            # Add blank lines for manual entry
            for i in range(5):
                doc.add_paragraph("_________________________________________________________________")
        
        doc.add_paragraph()  # Add spacing
        
        # 3. Communication Channels Section
        doc.add_heading("3. قنوات التواصل", level=2)
        
        # Internal Communications
        doc.add_heading("التواصل الداخلي:", level=3)
        internal_comms = form_data.get('internal_communications', [])
        if internal_comms:
            comm_table = doc.add_table(rows=1, cols=2)
            comm_table.style = 'Table Grid'
            comm_table.rows[0].cells[0].text = "جهات التواصل"
            comm_table.rows[0].cells[1].text = "الغرض من التواصل"
            
            for comm in internal_comms:
                if comm.get('entity') or comm.get('purpose'):
                    row = comm_table.add_row()
                    row.cells[0].text = comm.get('entity', '') or "_________________"
                    row.cells[1].text = comm.get('purpose', '') or "_________________"
        else:
            # Add blank lines for manual entry
            for i in range(3):
                doc.add_paragraph("جهات التواصل: _________________ الغرض: _________________")
        
        # External Communications
        doc.add_heading("التواصل الخارجي:", level=3)
        external_comms = form_data.get('external_communications', [])
        if external_comms:
            for comm in external_comms:
                if comm.get('entity') or comm.get('purpose'):
                    doc.add_paragraph(f"جهات التواصل: {comm.get('entity', '') or '_________________'} الغرض: {comm.get('purpose', '') or '_________________'}")
        else:
            # Add blank lines for manual entry
            for i in range(3):
                doc.add_paragraph("جهات التواصل: _________________ الغرض: _________________")
        
        doc.add_paragraph()  # Add spacing
        
        # 4. Job Standard Levels Section
        doc.add_heading("4. مستويات المهنة القياسية", level=2)
        job_levels = form_data.get('job_levels', [])
        if job_levels:
            level_table = doc.add_table(rows=1, cols=4)
            level_table.style = 'Table Grid'
            level_table.rows[0].cells[0].text = "المستوى"
            level_table.rows[0].cells[1].text = "الرمز"
            level_table.rows[0].cells[2].text = "الدور"
            level_table.rows[0].cells[3].text = "التقدم"
            
            for level in job_levels:
                if any(level.values()):
                    row = level_table.add_row()
                    row.cells[0].text = level.get('level', '') or "_________________"
                    row.cells[1].text = level.get('code', '') or "_________________"
                    row.cells[2].text = level.get('role', '') or "_________________"
                    row.cells[3].text = level.get('progression', '') or "_________________"
        else:
            # Add blank table for manual entry
            level_table = doc.add_table(rows=2, cols=4)
            level_table.style = 'Table Grid'
            level_table.rows[0].cells[0].text = "المستوى"
            level_table.rows[0].cells[1].text = "الرمز"
            level_table.rows[0].cells[2].text = "الدور"
            level_table.rows[0].cells[3].text = "التقدم"
            level_table.rows[1].cells[0].text = "_________________"
            level_table.rows[1].cells[1].text = "_________________"
            level_table.rows[1].cells[2].text = "_________________"
            level_table.rows[1].cells[3].text = "_________________"
        
        doc.add_paragraph()  # Add spacing
        
        # 5. Competencies Section
        doc.add_heading("5. الجدارات", level=2)
        
        # Behavioral Competencies
        doc.add_heading("الجدارات السلوكية:", level=3)
        behavioral_comp = form_data.get('behavioral_competencies', [])
        if behavioral_comp:
            for comp in behavioral_comp:
                if comp.get('name') or comp.get('level'):
                    doc.add_paragraph(f"• {comp.get('name', '') or '_________________'} - المستوى: {comp.get('level', '') or '_________________'}")
        else:
            # Add blank lines for manual entry
            for i in range(5):
                doc.add_paragraph(f"• _________________ - المستوى: _________________")
        
        # Core Competencies
        doc.add_heading("الجدارات الأساسية:", level=3)
        core_comp = form_data.get('core_competencies', [])
        if core_comp:
            for comp in core_comp:
                if comp.get('name') or comp.get('level'):
                    doc.add_paragraph(f"• {comp.get('name', '') or '_________________'} - المستوى: {comp.get('level', '') or '_________________'}")
        else:
            # Add blank lines for manual entry
            for i in range(5):
                doc.add_paragraph(f"• _________________ - المستوى: _________________")
        
        # Leadership Competencies
        doc.add_heading("الجدارات القيادية:", level=3)
        leadership_comp = form_data.get('leadership_competencies', [])
        if leadership_comp:
            for comp in leadership_comp:
                if comp.get('name') or comp.get('level'):
                    doc.add_paragraph(f"• {comp.get('name', '') or '_________________'} - المستوى: {comp.get('level', '') or '_________________'}")
        else:
            # Add blank lines for manual entry
            for i in range(5):
                doc.add_paragraph(f"• _________________ - المستوى: _________________")
        
        # Technical Competencies
        doc.add_heading("الجدارات التقنية:", level=3)
        technical_comp = form_data.get('technical_competencies', [])
        if technical_comp:
            for comp in technical_comp:
                if comp.get('name') or comp.get('level'):
                    doc.add_paragraph(f"• {comp.get('name', '') or '_________________'} - المستوى: {comp.get('level', '') or '_________________'}")
        else:
            # Add blank lines for manual entry
            for i in range(5):
                doc.add_paragraph(f"• _________________ - المستوى: _________________")
        
        # Page break for second part
        doc.add_page_break()
        
        # ===== PART B: نموذج الوصف الفعلي =====
        doc.add_heading("ب- نموذج الوصف الفعلي", level=1)
        doc.add_paragraph()
        
        # 1. Tasks Section
        doc.add_heading("1. المهام", level=2)
        
        # Leadership Tasks
        doc.add_heading("المهام القيادية/الإشرافية:", level=3)
        leadership_tasks = form_data.get('leadership_tasks', [])
        if leadership_tasks:
            for task in leadership_tasks:
                if task:
                    doc.add_paragraph(f"• {task}")
        else:
            # Add blank lines for manual entry
            for i in range(5):
                doc.add_paragraph("• _________________")
        
        # Specialized Tasks
        doc.add_heading("المهام التخصصية:", level=3)
        specialized_tasks = form_data.get('specialized_tasks', [])
        if specialized_tasks:
            for task in specialized_tasks:
                if task:
                    doc.add_paragraph(f"• {task}")
        else:
            # Add blank lines for manual entry
            for i in range(5):
                doc.add_paragraph("• _________________")
        
        # Other Tasks
        doc.add_heading("المهام الإضافية:", level=3)
        other_tasks = form_data.get('other_tasks', [])
        if other_tasks:
            for task in other_tasks:
                if task:
                    doc.add_paragraph(f"• {task}")
        else:
            # Add blank lines for manual entry
            for i in range(5):
                doc.add_paragraph("• _________________")
        
        doc.add_paragraph()  # Add spacing
        
        # 2. Competency Tables Section
        doc.add_heading("2. الجدارات السلوكية والفنية", level=2)
        
        # Behavioral Competencies Table
        doc.add_heading("الجدارات السلوكية:", level=3)
        behavioral_table = doc.add_table(rows=1, cols=3)
        behavioral_table.style = 'Table Grid'
        behavioral_table.rows[0].cells[0].text = "الرقم"
        behavioral_table.rows[0].cells[1].text = "اسم الجدارة"
        behavioral_table.rows[0].cells[2].text = "مستوى الإتقان"
        
        behavioral_data = form_data.get('behavioral_table', [])
        if behavioral_data:
            for comp in behavioral_data:
                if comp.get('name') or comp.get('level'):
                    row = behavioral_table.add_row()
                    row.cells[0].text = str(comp.get('number', '')) or "_________________"
                    row.cells[1].text = comp.get('name', '') or "_________________"
                    row.cells[2].text = comp.get('level', '') or "_________________"
        else:
            # Add blank rows for manual entry
            for i in range(5):
                row = behavioral_table.add_row()
                row.cells[0].text = str(i + 1)
                row.cells[1].text = "_________________"
                row.cells[2].text = "_________________"
        
        doc.add_paragraph()  # Add spacing
        
        # Technical Competencies Table
        doc.add_heading("الجدارات الفنية:", level=3)
        technical_table = doc.add_table(rows=1, cols=3)
        technical_table.style = 'Table Grid'
        technical_table.rows[0].cells[0].text = "الرقم"
        technical_table.rows[0].cells[1].text = "اسم الجدارة"
        technical_table.rows[0].cells[2].text = "مستوى الإتقان"
        
        technical_data = form_data.get('technical_table', [])
        if technical_data:
            for comp in technical_data:
                if comp.get('name') or comp.get('level'):
                    row = technical_table.add_row()
                    row.cells[0].text = str(comp.get('number', '')) or "_________________"
                    row.cells[1].text = comp.get('name', '') or "_________________"
                    row.cells[2].text = comp.get('level', '') or "_________________"
        else:
            # Add blank rows for manual entry
            for i in range(5):
                row = technical_table.add_row()
                row.cells[0].text = str(i + 1)
                row.cells[1].text = "_________________"
                row.cells[2].text = "_________________"
        
        doc.add_paragraph()  # Add spacing
        
        # 3. Performance Management Section
        doc.add_heading("3. إدارة الأداء المهني", level=2)
        
        # KPIs Table
        doc.add_heading("مؤشرات الأداء الرئيسية:", level=3)
        kpis_table = doc.add_table(rows=1, cols=3)
        kpis_table.style = 'Table Grid'
        kpis_table.rows[0].cells[0].text = "الرقم"
        kpis_table.rows[0].cells[1].text = "مؤشرات الأداء الرئيسية"
        kpis_table.rows[0].cells[2].text = "طريقة القياس"
        
        kpis_data = form_data.get('kpis', [])
        if kpis_data:
            for kpi in kpis_data:
                if kpi.get('metric') or kpi.get('measure'):
                    row = kpis_table.add_row()
                    row.cells[0].text = str(kpi.get('number', '')) or "_________________"
                    row.cells[1].text = kpi.get('metric', '') or "_________________"
                    row.cells[2].text = kpi.get('measure', '') or "_________________"
        else:
            # Add blank rows for manual entry
            for i in range(5):
                row = kpis_table.add_row()
                row.cells[0].text = str(i + 1)
                row.cells[1].text = "_________________"
                row.cells[2].text = "_________________"
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph("_" * 80)
        doc.add_paragraph("Powered by Professional Job Description System")
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
        
    except Exception as e:
        st.error(f"خطأ في إنشاء التقرير DOCX: {str(e)}")
        return None

def generate_pdf_report(form_data, ai_analysis=None):
    """Generate a professional PDF report from form data and AI analysis"""
    try:
        # Check if fonts are available and register them
        font_result = register_arabic_fonts()
        if font_result is True:
            # Use Noto Naskh Arabic fonts
            arabic_font = AR_FONT_REGULAR
            arabic_font_bold = AR_FONT_BOLD
        else:
            # Use fallback system font
            arabic_font = font_result
            arabic_font_bold = font_result
            
        # Show font status
        if font_result is True:
            st.success("✅ تم تسجيل الخطوط العربية بنجاح!")
        else:
            st.warning(f"⚠️ استخدام خط النظام: {font_result}")
            st.info("💡 للحصول على دعم كامل للعربية، قم بتثبيت الخطوط يدوياً")
        
        # Create a BytesIO buffer for the PDF
        buffer = io.BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles for Arabic text
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=arabic_font_bold,
            fontSize=24,
            alignment=TA_CENTER,
            spaceAfter=30,
            textColor=colors.darkblue
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontName=arabic_font,
            fontSize=14,
            alignment=TA_CENTER,
            spaceAfter=20,
            textColor=colors.gray
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=arabic_font_bold,
            fontSize=16,
            alignment=TA_RIGHT,
            spaceAfter=12,
            textColor=colors.darkblue,
            borderWidth=1,
            borderColor=colors.darkblue,
            borderPadding=5,
            backColor=colors.lightblue
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading3'],
            fontName=arabic_font_bold,
            fontSize=13,
            alignment=TA_RIGHT,
            spaceAfter=8,
            textColor=colors.black
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=arabic_font,
            fontSize=12,
            alignment=TA_RIGHT,
            spaceAfter=6
        )
        
        highlight_style = ParagraphStyle(
            'CustomHighlight',
            parent=styles['Normal'],
            fontName=arabic_font_bold,
            fontSize=12,
            alignment=TA_RIGHT,
            textColor=colors.darkred,
            spaceAfter=6
        )
        
        # Title
        story.append(Paragraph(A("نظام بطاقة الوصف المهني"), title_style))
        story.append(Paragraph("Professional Job Description Card System", subtitle_style))
        story.append(Spacer(1, 30))
        
        # Add timestamp
        from datetime import datetime
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        story.append(Paragraph(A(f"تاريخ الإنشاء: {current_time}"), normal_style))
        story.append(Spacer(1, 20))
        
        # Reference Data Section
        story.append(Paragraph(A("أ‌- البيانات المرجعية للمهنة"), heading_style))
        story.append(Spacer(1, 10))
        
        ref_data = form_data.get('ref_data', {})
        ref_table_data = [
            [A("المجال"), A("القيمة")],
            [A("المجموعة الرئيسية"), A(ref_data.get('main_group', ''))],
            [A("رمز المجموعة الرئيسية"), A(ref_data.get('main_group_code', ''))],
            [A("المجموعة الفرعية"), A(ref_data.get('sub_group', ''))],
            [A("رمز المجموعة الفرعية"), A(ref_data.get('sub_group_code', ''))],
            [A("المجموعة الثانوية"), A(ref_data.get('secondary_group', ''))],
            [A("رمز المجموعة الثانوية"), A(ref_data.get('secondary_group_code', ''))],
            [A("مجموعة الوحدات"), A(ref_data.get('unit_group', ''))],
            [A("رمز الوحدات"), A(ref_data.get('unit_group_code', ''))],
            [A("المهنة"), A(ref_data.get('job', ''))],
            [A("رمز المهنة"), A(ref_data.get('job_code', ''))],
            [A("موقع العمل"), A(ref_data.get('work_location', ''))],
            [A("المرتبة"), A(ref_data.get('grade', ''))]
        ]
        
        ref_table = Table(ref_table_data, colWidths=[2.5*inch, 3.5*inch])
        ref_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 15),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.darkblue),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.lightblue, colors.white]),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, -1), arabic_font)
        ]))
        story.append(ref_table)
        story.append(Spacer(1, 25))
        
        # Summary Section
        if form_data.get('summary'):
            story.append(Paragraph(A("ب‌- ملخص الوظيفة"), heading_style))
            story.append(Spacer(1, 10))
            
            # Add summary in a highlighted box
            summary_text = form_data.get('summary', '')
            if summary_text:
                story.append(Paragraph(A(f"الملخص: {summary_text}"), highlight_style))
            story.append(Spacer(1, 25))
        
        # Communications Section
        story.append(Paragraph(A("ج‌- قنوات التواصل"), heading_style))
        story.append(Spacer(1, 10))
        
        # Internal Communications
        story.append(Paragraph(A("التواصل الداخلي:"), subheading_style))
        internal_comms = form_data.get('internal_communications', [])
        if internal_comms and any(any(comm.values()) for comm in internal_comms):
            for i, comm in enumerate(internal_comms, 1):
                if any(comm.values()):
                    story.append(Paragraph(A(f"• {comm.get('entity', '')} - {comm.get('purpose', '')}"), normal_style))
        else:
            story.append(Paragraph(A("لا توجد بيانات"), normal_style))
        
        story.append(Spacer(1, 15))
        
        # External Communications
        story.append(Paragraph(A("التواصل الخارجي:"), subheading_style))
        external_comms = form_data.get('external_communications', [])
        if external_comms and any(any(comm.values()) for comm in external_comms):
            for i, comm in enumerate(external_comms, 1):
                if any(comm.values()):
                    story.append(Paragraph(A(f"• {comm.get('entity', '')} - {comm.get('purpose', '')}"), normal_style))
        else:
            story.append(Paragraph(A("لا توجد بيانات"), normal_style))
        
        story.append(Spacer(1, 25))
        
        # Job Levels Section
        story.append(Paragraph(A("د‌- مستويات الوظيفة"), heading_style))
        story.append(Spacer(1, 10))
        
        job_levels = form_data.get('job_levels', [])
        if job_levels and any(any(level.values()) for level in job_levels):
            level_table_data = [[A("المستوى"), A("الرمز"), A("الدور"), A("التقدم")]]
            for level in job_levels:
                if any(level.values()):
                    level_table_data.append([
                        A(level.get('level', '')),
                        A(level.get('code', '')),
                        A(level.get('role', '')),
                        A(level.get('progression', ''))
                    ])
            
            level_table = Table(level_table_data, colWidths=[1.5*inch, 1*inch, 2*inch, 1.5*inch])
            level_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.lightgreen),
            ('GRID', (0, 0), (-1, -1), 1, colors.darkgreen),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.lightgreen, colors.white]),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, -1), arabic_font)
            ]))
            story.append(level_table)
        else:
            story.append(Paragraph(A("لا توجد بيانات"), normal_style))
        
        story.append(Spacer(1, 25))
        
        # Competencies Section
        story.append(Paragraph(A("هـ- الكفاءات المطلوبة"), heading_style))
        story.append(Spacer(1, 10))
        
        # Behavioral Competencies
        story.append(Paragraph(A("الكفاءات السلوكية:"), subheading_style))
        behavioral_comps = form_data.get('behavioral_competencies', [])
        if behavioral_comps and any(any(comp.values()) for comp in behavioral_comps):
            for i, comp in enumerate(behavioral_comps, 1):
                if any(comp.values()):
                    story.append(Paragraph(A(f"• {comp.get('name', '')} - المستوى: {comp.get('level', '')}"), normal_style))
        else:
            story.append(Paragraph(A("لا توجد بيانات"), normal_style))
        
        story.append(Spacer(1, 15))
        
        # Core Competencies
        story.append(Paragraph(A("الكفاءات الأساسية:"), subheading_style))
        core_comps = form_data.get('core_competencies', [])
        if core_comps and any(any(comp.values()) for comp in core_comps):
            for i, comp in enumerate(core_comps, 1):
                if any(comp.values()):
                    story.append(Paragraph(A(f"• {comp.get('name', '')} - المستوى: {comp.get('level', '')}"), normal_style))
        else:
            story.append(Paragraph(A("لا توجد بيانات"), normal_style))
        
        story.append(Spacer(1, 15))
        
        # Leadership Competencies
        story.append(Paragraph(A("الكفاءات القيادية:"), subheading_style))
        leadership_comps = form_data.get('leadership_competencies', [])
        if leadership_comps and any(any(comp.values()) for comp in leadership_comps):
            for i, comp in enumerate(leadership_comps, 1):
                if any(comp.values()):
                    story.append(Paragraph(A(f"• {comp.get('name', '')} - المستوى: {comp.get('level', '')}"), normal_style))
        else:
            story.append(Paragraph(A("لا توجد بيانات"), normal_style))
        
        story.append(Spacer(1, 15))
        
        # Technical Competencies
        story.append(Paragraph(A("الكفاءات التقنية:"), subheading_style))
        technical_comps = form_data.get('technical_competencies', [])
        if technical_comps and any(any(comp.values()) for comp in technical_comps):
            for i, comp in enumerate(technical_comps, 1):
                if any(comp.values()):
                    story.append(Paragraph(A(f"• {comp.get('name', '')} - المستوى: {comp.get('level', '')}"), normal_style))
        else:
            story.append(Paragraph(A("لا توجد بيانات"), normal_style))
        
        story.append(Spacer(1, 25))
        
        # Tasks Section
        story.append(Paragraph(A("و‌- المهام"), heading_style))
        story.append(Spacer(1, 10))
        
        # Leadership Tasks
        story.append(Paragraph(A("المهام القيادية:"), subheading_style))
        leadership_tasks = form_data.get('leadership_tasks', [])
        if leadership_tasks and any(task for task in leadership_tasks):
            for i, task in enumerate(leadership_tasks, 1):
                if task:
                    story.append(Paragraph(A(f"{i}. {task}"), normal_style))
        else:
            story.append(Paragraph(A("لا توجد بيانات"), normal_style))
        
        story.append(Spacer(1, 15))
        
        # Specialized Tasks
        story.append(Paragraph(A("المهام المتخصصة:"), subheading_style))
        specialized_tasks = form_data.get('specialized_tasks', [])
        if specialized_tasks and any(task for task in specialized_tasks):
            for i, task in enumerate(specialized_tasks, 1):
                if task:
                    story.append(Paragraph(A(f"{i}. {task}"), normal_style))
        else:
            story.append(Paragraph(A("لا توجد بيانات"), normal_style))
        
        story.append(Spacer(1, 15))
        
        # Other Tasks
        story.append(Paragraph(A("المهام الأخرى:"), subheading_style))
        other_tasks = form_data.get('other_tasks', [])
        if other_tasks and any(task for task in other_tasks):
            for i, task in enumerate(other_tasks, 1):
                if task:
                    story.append(Paragraph(A(f"{i}. {task}"), normal_style))
        else:
            story.append(Paragraph(A("لا توجد بيانات"), normal_style))
        
        story.append(Spacer(1, 25))
        
        # KPIs Section
        story.append(Paragraph(A("ز‌- مؤشرات الأداء الرئيسية"), heading_style))
        story.append(Spacer(1, 10))
        
        kpis = form_data.get('kpis', [])
        if kpis and any(any(kpi.values()) for kpi in kpis):
            kpi_table_data = [[A("الرقم"), A("المؤشر"), A("القياس")]]
            for kpi in kpis:
                if any(kpi.values()):
                    kpi_table_data.append([
                        str(kpi.get('number', '')),
                        A(kpi.get('metric', '')),
                        A(kpi.get('measure', ''))
                    ])
            
            kpi_table = Table(kpi_table_data, colWidths=[0.5*inch, 2.5*inch, 2*inch])
            kpi_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkred),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.lightcoral),
            ('GRID', (0, 0), (-1, -1), 1, colors.darkred),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.lightcoral, colors.white]),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, -1), arabic_font)
            ]))
            story.append(kpi_table)
        else:
            story.append(Paragraph(process_arabic_text("لا توجد بيانات"), normal_style))
        
        # AI Analysis Section (if available)
        if ai_analysis:
            story.append(PageBreak())
            story.append(Paragraph(A("تحليل الذكاء الاصطناعي"), title_style))
            story.append(Spacer(1, 20))
            
            # Show AI analysis in a formatted way
            try:
                ai_data = json.loads(ai_analysis)
                story.append(Paragraph(A("ملخص التحليل:"), heading_style))
                story.append(Spacer(1, 10))
                
                # Show key insights from AI
                if 'summary' in ai_data and ai_data['summary']:
                    story.append(Paragraph(A(f"الملخص: {ai_data['summary']}"), normal_style))
                    story.append(Spacer(1, 10))
                
                # Show extracted competencies count
                total_competencies = 0
                for comp_type in ['behavioral_competencies', 'core_competencies', 'leadership_competencies', 'technical_competencies']:
                    if comp_type in ai_data:
                        count = len([c for c in ai_data[comp_type] if any(c.values())])
                        total_competencies += count
                
                if total_competencies > 0:
                    story.append(Paragraph(A(f"إجمالي الكفاءات المستخرجة: {total_competencies}"), highlight_style))
                    story.append(Spacer(1, 10))
                
                # Show tasks count
                total_tasks = 0
                for task_type in ['leadership_tasks', 'specialized_tasks', 'other_tasks']:
                    if task_type in ai_data:
                        count = len([t for t in ai_data[task_type] if t])
                        total_tasks += count
                
                if total_tasks > 0:
                    story.append(Paragraph(A(f"إجمالي المهام المستخرجة: {total_tasks}"), highlight_style))
                
            except json.JSONDecodeError:
                story.append(Paragraph(A("تحليل نصي:"), heading_style))
                story.append(Paragraph(A(ai_analysis[:1000] + "..." if len(ai_analysis) > 1000 else ai_analysis), normal_style))
        
        # Add footer
        story.append(Spacer(1, 30))
        story.append(Paragraph("─" * 50, normal_style))
        story.append(Spacer(1, 10))
        story.append(Paragraph(A("تم إنشاء هذا التقرير بواسطة نظام بطاقة الوصف المهني"), normal_style))
        story.append(Paragraph("Powered by AI-Powered Job Description System", normal_style))
        
        # Build the PDF
        doc.build(story)
        
        # Get the PDF content
        pdf_content = buffer.getvalue()
        buffer.close()
        
        return pdf_content
        
    except Exception as e:
        st.error(f"❌ خطأ في إنشاء PDF: {str(e)}")
        return None

def render_reference_data():
    """Render the reference data section"""
    st.markdown('<div class="section-header">أ‌- نموذج بطاقة الوصف المهني</div>', unsafe_allow_html=True)
    st.markdown('<div class="subsection-header">1- البيانات المرجعية للمهنة</div>', unsafe_allow_html=True)
    
    with st.container():
        col1, col2 = st.columns(2)
        
        with col1:
            st.session_state.form_data['ref_data']['main_group'] = st.text_input(
                "المجموعة الرئيسية",
                value=st.session_state.form_data['ref_data']['main_group'],
                key="main_group"
            )
            
            st.session_state.form_data['ref_data']['sub_group'] = st.text_input(
                "المجموعة الفرعية",
                value=st.session_state.form_data['ref_data']['sub_group'],
                key="sub_group"
            )
            
            st.session_state.form_data['ref_data']['secondary_group'] = st.text_input(
                "المجموعة الثانوية",
                value=st.session_state.form_data['ref_data']['secondary_group'],
                key="secondary_group"
            )
            
            st.session_state.form_data['ref_data']['unit_group'] = st.text_input(
                "مجموعة الوحدات",
                value=st.session_state.form_data['ref_data']['unit_group'],
                key="unit_group"
            )
            
            st.session_state.form_data['ref_data']['job'] = st.text_input(
                "المهنة *",
                value=st.session_state.form_data['ref_data']['job'],
                key="job"
            )
            
            st.session_state.form_data['ref_data']['work_location'] = st.text_input(
                "موقع العمل *",
                value=st.session_state.form_data['ref_data']['work_location'],
                key="work_location"
            )
        
        with col2:
            st.session_state.form_data['ref_data']['main_group_code'] = st.text_input(
                "رمز المجموعة الرئيسية",
                value=st.session_state.form_data['ref_data']['main_group_code'],
                key="main_group_code"
            )
            
            st.session_state.form_data['ref_data']['sub_group_code'] = st.text_input(
                "رمز المجموعة الفرعية",
                value=st.session_state.form_data['ref_data']['sub_group_code'],
                key="sub_group_code"
            )
            
            st.session_state.form_data['ref_data']['secondary_group_code'] = st.text_input(
                "رمز المجموعة الثانوية",
                value=st.session_state.form_data['ref_data']['secondary_group_code'],
                key="secondary_group_code"
            )
            
            st.session_state.form_data['ref_data']['unit_group_code'] = st.text_input(
                "رمز الوحدات",
                value=st.session_state.form_data['ref_data']['unit_group_code'],
                key="unit_group_code"
            )
            
            st.session_state.form_data['ref_data']['job_code'] = st.text_input(
                "رمز المهنة",
                value=st.session_state.form_data['ref_data']['job_code'],
                key="job_code"
            )
            
            st.session_state.form_data['ref_data']['grade'] = st.text_input(
                "المرتبة",
                value=st.session_state.form_data['ref_data']['grade'],
                key="grade"
            )

def render_summary():
    """Render the job summary section"""
    st.markdown('<div class="subsection-header">2- الملخص العام للمهنة</div>', unsafe_allow_html=True)
    
    st.session_state.form_data['summary'] = st.text_area(
        "الملخص العام للمهنة",
        value=st.session_state.form_data['summary'],
        height=120,
        key="summary"
    )

def render_communication_channels():
    """Render the communication channels section"""
    st.markdown('<div class="subsection-header">3- قنوات التواصل</div>', unsafe_allow_html=True)
    
    # Internal communications
    st.markdown("**الجهات الداخلية:**")
    for i, comm in enumerate(st.session_state.form_data['internal_communications']):
        with st.container():
            col1, col2, col3 = st.columns([3, 3, 1])
            with col1:
                comm['entity'] = st.text_input(
                    "الجهة",
                    value=comm['entity'],
                    key=f"internal_entity_{i}"
                )
            with col2:
                comm['purpose'] = st.text_input(
                    "الغرض من التواصل",
                    value=comm['purpose'],
                    key=f"internal_purpose_{i}"
                )
            with col3:
                if st.button("حذف", key=f"remove_internal_{i}", type="secondary"):
                    remove_row(st.session_state.form_data['internal_communications'], i)
                    break
    
    if st.button("+ إضافة جهة داخلية", key="add_internal", type="primary"):
        add_row(st.session_state.form_data['internal_communications'], {'entity': '', 'purpose': ''})
    
    st.markdown("---")
    
    # External communications
    st.markdown("**الجهات الخارجية:**")
    for i, comm in enumerate(st.session_state.form_data['external_communications']):
        with st.container():
            col1, col2, col3 = st.columns([3, 3, 1])
            with col1:
                comm['entity'] = st.text_input(
                    "الجهة",
                    value=comm['entity'],
                    key=f"external_entity_{i}"
                )
            with col2:
                comm['purpose'] = st.text_input(
                    "الغرض من التواصل",
                    value=comm['purpose'],
                    key=f"external_purpose_{i}"
                )
            with col3:
                if st.button("حذف", key=f"remove_external_{i}", type="secondary"):
                    remove_row(st.session_state.form_data['external_communications'], i)
                    break
    
    if st.button("+ إضافة جهة خارجية", key="add_external", type="primary"):
        add_row(st.session_state.form_data['external_communications'], {'entity': '', 'purpose': ''})

def render_job_levels():
    """Render the job levels section"""
    st.markdown('<div class="subsection-header">4- مستويات المهنة القياسية</div>', unsafe_allow_html=True)
    
    for i, level in enumerate(st.session_state.form_data['job_levels']):
        with st.container():
            col1, col2, col3, col4, col5 = st.columns([2, 2, 2, 2, 1])
            with col1:
                level['level'] = st.text_input(
                    "مستوى المهنة القياسي",
                    value=level['level'],
                    key=f"job_level_{i}"
                )
            with col2:
                level['code'] = st.text_input(
                    "رمز المستوى المهني",
                    value=level['code'],
                    key=f"job_code_{i}"
                )
            with col3:
                level['role'] = st.text_input(
                    "الدور المهني",
                    value=level['role'],
                    key=f"job_role_{i}"
                )
            with col4:
                level['progression'] = st.text_input(
                    "التدرج المهني",
                    value=level['progression'],
                    key=f"job_progression_{i}"
                )
            with col5:
                if st.button("حذف", key=f"remove_level_{i}", type="secondary"):
                    remove_row(st.session_state.form_data['job_levels'], i)
                    break
    
    if st.button("+ إضافة مستوى", key="add_level", type="primary"):
        add_row(st.session_state.form_data['job_levels'], {'level': '', 'code': '', 'role': '', 'progression': ''})

def render_competencies():
    """Render the competencies section"""
    st.markdown('<div class="subsection-header">5- الجدارات</div>', unsafe_allow_html=True)
    
    # Behavioral competencies
    st.markdown("**الجدارات السلوكية:**")
    for i, comp in enumerate(st.session_state.form_data['behavioral_competencies']):
        with st.container():
            col1, col2, col3 = st.columns([3, 2, 1])
            with col1:
                comp['name'] = st.text_input(
                    "الجدارة",
                    value=comp['name'],
                    key=f"behavioral_name_{i}"
                )
            with col2:
                comp['level'] = st.text_input(
                    "مستوى الإتقان",
                    value=comp['level'],
                    key=f"behavioral_level_{i}"
                )
            with col3:
                if st.button("حذف", key=f"remove_behavioral_{i}", type="secondary"):
                    remove_row(st.session_state.form_data['behavioral_competencies'], i)
                    break
    
    if st.button("+ إضافة جدارة سلوكية", key="add_behavioral", type="primary"):
        add_row(st.session_state.form_data['behavioral_competencies'], {'name': '', 'level': ''})
    
    st.markdown("---")
    
    # Core competencies
    st.markdown("**الجدارات الأساسية:**")
    for i, comp in enumerate(st.session_state.form_data['core_competencies']):
        with st.container():
            col1, col2, col3 = st.columns([3, 2, 1])
            with col1:
                comp['name'] = st.text_input(
                    "الجدارة",
                    value=comp['name'],
                    key=f"core_name_{i}"
                )
            with col2:
                comp['level'] = st.text_input(
                    "مستوى الإتقان",
                    value=comp['level'],
                    key=f"core_level_{i}"
                )
            with col3:
                if st.button("حذف", key=f"remove_core_{i}", type="secondary"):
                    remove_row(st.session_state.form_data['core_competencies'], i)
                    break
    
    if st.button("+ إضافة جدارة أساسية", key="add_core", type="primary"):
        add_row(st.session_state.form_data['core_competencies'], {'name': '', 'level': ''})
    
    st.markdown("---")
    
    # Leadership competencies
    st.markdown("**الجدارات القيادية:**")
    for i, comp in enumerate(st.session_state.form_data['leadership_competencies']):
        with st.container():
            col1, col2, col3 = st.columns([3, 2, 1])
            with col1:
                comp['name'] = st.text_input(
                    "الجدارة",
                    value=comp['name'],
                    key=f"leadership_name_{i}"
                )
            with col2:
                comp['level'] = st.text_input(
                    "مستوى الإتقان",
                    value=comp['level'],
                    key=f"leadership_level_{i}"
                )
            with col3:
                if st.button("حذف", key=f"remove_leadership_{i}", type="secondary"):
                    remove_row(st.session_state.form_data['leadership_competencies'], i)
                    break
    
    if st.button("+ إضافة جدارة قيادية", key="add_leadership", type="primary"):
        add_row(st.session_state.form_data['leadership_competencies'], {'name': '', 'level': ''})
    
    st.markdown("---")
    
    # Technical competencies
    st.markdown("**الجدارات الفنية:**")
    for i, comp in enumerate(st.session_state.form_data['technical_competencies']):
        with st.container():
            col1, col2, col3 = st.columns([3, 2, 1])
            with col1:
                comp['name'] = st.text_input(
                    "الجدارة",
                    value=comp['name'],
                    key=f"technical_name_{i}"
                )
            with col2:
                comp['level'] = st.text_input(
                    "مستوى الإتقان",
                    value=comp['level'],
                    key=f"technical_level_{i}"
                )
            with col3:
                if st.button("حذف", key=f"remove_technical_{i}", type="secondary"):
                    remove_row(st.session_state.form_data['technical_competencies'], i)
                    break
    
    if st.button("+ إضافة جدارة فنية", key="add_technical", type="primary"):
        add_row(st.session_state.form_data['technical_competencies'], {'name': '', 'level': ''})

def render_actual_description():
    """Render the actual description section"""
    st.markdown('<div class="section-header">ب‌- نموذج الوصف الفعلي</div>', unsafe_allow_html=True)
    
    # Tasks section
    st.markdown('<div class="subsection-header">1- المهام</div>', unsafe_allow_html=True)
    
    # Leadership tasks
    st.markdown("**المهام القيادية/الإشرافية:**")
    for i, task in enumerate(st.session_state.form_data['leadership_tasks']):
        with st.container():
            col1, col2 = st.columns([4, 1])
            with col1:
                st.session_state.form_data['leadership_tasks'][i] = st.text_input(
                    "المهمة",
                    value=task,
                    key=f"leadership_task_{i}"
                )
            with col2:
                if st.button("حذف", key=f"remove_leadership_task_{i}", type="secondary"):
                    st.session_state.form_data['leadership_tasks'].pop(i)
                    st.rerun()
                    break
    
    if st.button("+ إضافة مهمة قيادية", key="add_leadership_task", type="primary"):
        st.session_state.form_data['leadership_tasks'].append('')
        st.rerun()
    
    st.markdown("---")
    
    # Specialized tasks
    st.markdown("**المهام التخصصية:**")
    for i, task in enumerate(st.session_state.form_data['specialized_tasks']):
        with st.container():
            col1, col2 = st.columns([4, 1])
            with col1:
                st.session_state.form_data['specialized_tasks'][i] = st.text_input(
                    "المهمة",
                    value=task,
                    key=f"specialized_task_{i}"
                )
            with col2:
                if st.button("حذف", key=f"remove_specialized_task_{i}", type="secondary"):
                    st.session_state.form_data['specialized_tasks'].pop(i)
                    st.rerun()
                    break
    
    if st.button("+ إضافة مهمة تخصصية", key="add_specialized_task", type="primary"):
        st.session_state.form_data['specialized_tasks'].append('')
        st.rerun()
    
    st.markdown("---")
    
    # Other tasks
    st.markdown("**مهام أخرى إضافية:**")
    for i, task in enumerate(st.session_state.form_data['other_tasks']):
        with st.container():
            col1, col2 = st.columns([4, 1])
            with col1:
                st.session_state.form_data['other_tasks'][i] = st.text_input(
                    "المهمة",
                    value=task,
                    key=f"other_task_{i}"
                )
            with col2:
                if st.button("حذف", key=f"remove_other_task_{i}", type="secondary"):
                    st.session_state.form_data['other_tasks'].pop(i)
                    st.rerun()
                    break
    
    if st.button("+ إضافة مهمة أخرى", key="add_other_task", type="primary"):
        st.session_state.form_data['other_tasks'].append('')
        st.rerun()

def render_competencies_tables():
    """Render the competencies tables section"""
    st.markdown('<div class="subsection-header">2- الجدارات السلوكية والفنية</div>', unsafe_allow_html=True)
    
    # Behavioral competencies table
    st.markdown("**الجدارات السلوكية:**")
    for i, comp in enumerate(st.session_state.form_data['behavioral_table']):
        with st.container():
            col1, col2, col3, col4 = st.columns([1, 3, 2, 1])
            with col1:
                comp['number'] = i + 1
                st.text_input(
                    "الرقم",
                    value=comp['number'],
                    key=f"behavioral_table_number_{i}",
                    disabled=True
                )
            with col2:
                comp['name'] = st.text_input(
                    "الجدارة",
                    value=comp['name'],
                    key=f"behavioral_table_name_{i}"
                )
            with col3:
                comp['level'] = st.text_input(
                    "مستوى الإتقان",
                    value=comp['level'],
                    key=f"behavioral_table_level_{i}"
                )
            with col4:
                if st.button("حذف", key=f"remove_behavioral_table_{i}", type="secondary"):
                    remove_row(st.session_state.form_data['behavioral_table'], i)
                    break
    
    if st.button("+ إضافة صف سلوكي", key="add_behavioral_table", type="primary"):
        new_number = len(st.session_state.form_data['behavioral_table']) + 1
        add_row(st.session_state.form_data['behavioral_table'], {'number': new_number, 'name': '', 'level': ''})
    
    st.markdown("---")
    
    # Technical competencies table
    st.markdown("**الجدارات الفنية:**")
    for i, comp in enumerate(st.session_state.form_data['technical_table']):
        with st.container():
            col1, col2, col3, col4 = st.columns([1, 3, 2, 1])
            with col1:
                comp['number'] = i + 1
                st.text_input(
                    "الرقم",
                    value=comp['number'],
                    key=f"technical_table_number_{i}",
                    disabled=True
                )
            with col2:
                comp['name'] = st.text_input(
                    "الجدارة",
                    value=comp['name'],
                    key=f"technical_table_name_{i}"
                )
            with col3:
                comp['level'] = st.text_input(
                    "مستوى الإتقان",
                    value=comp['level'],
                    key=f"technical_table_level_{i}"
                )
            with col4:
                if st.button("حذف", key=f"remove_technical_table_{i}", type="secondary"):
                    remove_row(st.session_state.form_data['technical_table'], i)
                    break
    
    if st.button("+ إضافة صف فني", key="add_technical_table", type="primary"):
        new_number = len(st.session_state.form_data['technical_table']) + 1
        add_row(st.session_state.form_data['technical_table'], {'number': new_number, 'name': '', 'level': ''})

def render_kpis():
    """Render the KPIs section"""
    st.markdown('<div class="subsection-header">3- إدارة الأداء المهني</div>', unsafe_allow_html=True)
    
    for i, kpi in enumerate(st.session_state.form_data['kpis']):
        with st.container():
            col1, col2, col3, col4 = st.columns([1, 3, 3, 1])
            with col1:
                kpi['number'] = i + 1
                st.text_input(
                    "الرقم",
                    value=kpi['number'],
                    key=f"kpi_number_{i}",
                    disabled=True
                )
            with col2:
                kpi['metric'] = st.text_input(
                    "مؤشرات الأداء الرئيسية",
                    value=kpi['metric'],
                    key=f"kpi_metric_{i}"
                )
            with col3:
                kpi['measure'] = st.text_input(
                    "طريقة القياس",
                    value=kpi['measure'],
                    key=f"kpi_measure_{i}"
                )
            with col4:
                if st.button("حذف", key=f"remove_kpi_{i}", type="secondary"):
                    remove_row(st.session_state.form_data['kpis'], i)
                    break
    
    if st.button("+ إضافة مؤشر أداء", key="add_kpi", type="primary"):
        new_number = len(st.session_state.form_data['kpis']) + 1
        add_row(st.session_state.form_data['kpis'], {'number': new_number, 'metric': '', 'measure': ''})

def validate_form() -> tuple[bool, List[str]]:
    """Validate the form and return validation status and errors"""
    errors = []
    
    # Required fields validation
    if not st.session_state.form_data['ref_data']['job'].strip():
        errors.append("حقل 'المهنة' مطلوب")
    
    if not st.session_state.form_data['ref_data']['work_location'].strip():
        errors.append("حقل 'موقع العمل' مطلوب")
    
    # Communication validation
    for i, comm in enumerate(st.session_state.form_data['internal_communications']):
        if comm['entity'].strip() and not comm['purpose'].strip():
            errors.append(f"جهة التواصل الداخلية {i+1}: يجب تحديد الغرض من التواصل")
    
    for i, comm in enumerate(st.session_state.form_data['external_communications']):
        if comm['entity'].strip() and not comm['purpose'].strip():
            errors.append(f"جهة التواصل الخارجية {i+1}: يجب تحديد الغرض من التواصل")
    
    return len(errors) == 0, errors

def generate_json_output() -> str:
    """Generate the final JSON output matching the schema"""
    output = {
        "ref": {
            "main_group": st.session_state.form_data['ref_data']['main_group'],
            "main_group_code": st.session_state.form_data['ref_data']['main_group_code'],
            "sub_group": st.session_state.form_data['ref_data']['sub_group'],
            "sub_group_code": st.session_state.form_data['ref_data']['sub_group_code'],
            "secondary_group": st.session_state.form_data['ref_data']['secondary_group'],
            "secondary_group_code": st.session_state.form_data['ref_data']['secondary_group_code'],
            "unit_group": st.session_state.form_data['ref_data']['unit_group'],
            "unit_group_code": st.session_state.form_data['ref_data']['unit_group_code'],
            "job": st.session_state.form_data['ref_data']['job'],
            "job_code": st.session_state.form_data['ref_data']['job_code'],
            "work_location": st.session_state.form_data['ref_data']['work_location'],
            "grade": st.session_state.form_data['ref_data']['grade']
        },
        "summary": st.session_state.form_data['summary'],
        "comm": {
            "internal": st.session_state.form_data['internal_communications'],
            "external": st.session_state.form_data['external_communications']
        },
        "levels": st.session_state.form_data['job_levels'],
        "comp": {
            "behavioral": st.session_state.form_data['behavioral_competencies'],
            "core": st.session_state.form_data['core_competencies'],
            "lead": st.session_state.form_data['leadership_competencies'],
            "tech": st.session_state.form_data['technical_competencies']
        },
        "tasks": {
            "lead": st.session_state.form_data['leadership_tasks'],
            "spec": st.session_state.form_data['specialized_tasks'],
            "other": st.session_state.form_data['other_tasks']
        },
        "beh": [{"name": comp['name'], "level": comp['level']} for comp in st.session_state.form_data['behavioral_table']],
        "tech": [{"name": comp['name'], "level": comp['level']} for comp in st.session_state.form_data['technical_table']],
        "kpis": st.session_state.form_data['kpis']
    }
    
    return json.dumps(output, ensure_ascii=False, indent=2)

def main():
    """Main application function"""
    # Initialize session state
    initialize_session_state()
    
    # Main header
    st.markdown('<div class="form-header">نظام بطاقة الوصف المهني</div>', unsafe_allow_html=True)
    
    # File Upload and AI Analysis Section
    # Simple text input section
    st.markdown('<div class="section-header">إدخال النص</div>', unsafe_allow_html=True)
    
    manual_text = st.text_area(
        "أدخل نص الوصف الوظيفي هنا:",
        height=150,
        placeholder="أدخل نص الوصف الوظيفي هنا...",
        help="يمكنك نسخ ولصق نص الوصف الوظيفي مباشرة هنا"
    )
    
    st.markdown("---")
    
    st.markdown("---")
    
    # Form sections
    render_reference_data()
    render_summary()
    render_communication_channels()
    render_job_levels()
    render_competencies()
    render_actual_description()
    render_competencies_tables()
    render_kpis()
    
    # Submit section
    st.markdown("---")
    st.markdown('<div class="section-header">حفظ وتصدير البيانات</div>', unsafe_allow_html=True)
    
    # Form validation and DOCX generation
    if st.button("إنشاء تقرير DOCX احترافي", key="generate_docx_main", type="primary", use_container_width=True):
        is_valid, errors = validate_form()
        
        if is_valid:
            st.success("تم التحقق من صحة البيانات بنجاح!")
            
            with st.spinner("جاري إنشاء التقرير DOCX..."):
                # Get AI analysis from session state if available
                ai_analysis = st.session_state.get('last_ai_analysis', None)
                
                # Generate DOCX
                docx_content = generate_docx_report(st.session_state.form_data, ai_analysis)
                
                if docx_content:
                    # Create filename with timestamp
                    from datetime import datetime
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"بطاقة_الوصف_المهني_{timestamp}.docx"
                    
                    # Download button
                    st.download_button(
                        label="تحميل التقرير DOCX",
                        data=docx_content,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    st.success(f"تم إنشاء التقرير DOCX بنجاح! يمكنك تحميله الآن.")
                    
                    # Show DOCX preview info
                    st.info("التقرير يتضمن:")
                    preview_items = []
                    if st.session_state.form_data.get('ref_data', {}).get('job'):
                        preview_items.append("• البيانات المرجعية للمهنة")
                    if st.session_state.form_data.get('summary'):
                        preview_items.append("• ملخص الوظيفة")
                    if any(st.session_state.form_data.get('internal_communications', [])):
                        preview_items.append("• قنوات التواصل")
                    if any(st.session_state.form_data.get('behavioral_competencies', [])):
                        preview_items.append("• الكفاءات المطلوبة")
                    if any(st.session_state.form_data.get('leadership_tasks', [])):
                        preview_items.append("• المهام والمسؤوليات")
                    if any(st.session_state.form_data.get('kpis', [])):
                        preview_items.append("• مؤشرات الأداء")
                    
                    for item in preview_items:
                        st.write(item)
                    
                else:
                    st.error("فشل في إنشاء التقرير DOCX")
        else:
            st.error("يوجد أخطاء في البيانات:")
            for error in errors:
                st.error(f"• {error}")
    
    # Additional options in columns
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("إعادة تعيين", key="reset_form", type="secondary", use_container_width=True):
            st.session_state.form_data = {
                'ref_data': {
                    'main_group': '', 'main_group_code': '', 'sub_group': '', 'sub_group_code': '',
                    'secondary_group': '', 'secondary_group_code': '', 'unit_group': '', 'unit_group_code': '',
                    'job': '', 'job_code': '', 'work_location': '', 'grade': ''
                },
                'summary': '',
                'internal_communications': [{'entity': '', 'purpose': ''}],
                'external_communications': [{'entity': '', 'purpose': ''}],
                'job_levels': [{'level': '', 'code': '', 'role': '', 'progression': ''}],
                'behavioral_competencies': [{'name': '', 'level': ''}],
                'core_competencies': [{'name': '', 'level': ''}],
                'leadership_competencies': [{'name': '', 'level': ''}],
                'technical_competencies': [{'name': '', 'level': ''}],
                'leadership_tasks': [''],
                'specialized_tasks': [''],
                'other_tasks': [''],
                'behavioral_table': [{'number': 1, 'name': '', 'level': ''}],
                'technical_table': [{'number': 1, 'name': '', 'level': ''}],
                'kpis': [{'number': 1, 'metric': '', 'measure': ''}]
            }
            st.rerun()
    
    with col2:
        if st.button("معاينة البيانات", key="preview_data", type="secondary", use_container_width=True):
            is_valid, errors = validate_form()
            if is_valid:
                st.success("تم التحقق من صحة البيانات بنجاح!")
                st.info("يمكنك الآن إنشاء تقرير DOCX")
            else:
                st.error("يوجد أخطاء في البيانات:")
                for error in errors:
                    st.error(f"• {error}")
    
    with col3:
        st.info("استخدم زر 'إنشاء تقرير DOCX احترافي' أعلاه لإنشاء التقرير")

if __name__ == "__main__":
    main()