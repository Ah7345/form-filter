from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tc

def set_cell_shading(cell, hex_color):
    """Apply table cell background color using OXML"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="000000", size=6):
    """Apply single-line borders to all sides of a cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Remove existing borders
    for border in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(border)
    
    # Add new borders
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    
    tcPr.append(tcBorders)

def set_col_widths(table, widths_in_cm):
    """Set exact column widths in centimeters"""
    for i, width in enumerate(widths_in_cm):
        for cell in table.columns[i].cells:
            cell.width = Cm(width)

def arabic(p):
    """Force paragraph RTL + right alignment for Arabic text"""
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True

def merge_vertically(table, col_idx, row_start, row_end):
    """Merge cells vertically in a table"""
    for row_idx in range(row_start, row_end + 1):
        if row_idx == row_start:
            # First cell - add vMerge start
            cell = table.cell(row_idx, col_idx)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vMerge = OxmlElement('w:vMerge')
            vMerge.set(qn('w:val'), 'restart')
            tcPr.append(vMerge)
        else:
            # Subsequent cells - add vMerge continue
            cell = table.cell(row_idx, col_idx)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vMerge = OxmlElement('w:vMerge')
            vMerge.set(qn('w:val'), 'continue')
            tcPr.append(vMerge)

def create_header_band(doc, text):
    """Create a header band table with the specified text"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    
    # Set cell properties
    set_cell_shading(cell, "D9D9D9")
    set_cell_borders(cell)
    
    # Add text
    p = cell.paragraphs[0]
    p.text = text
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_to_left = True
    
    # Make text bold and 12pt
    for run in p.runs:
        run.font.bold = True
        run.font.size = Cm(0.42)  # 12pt
    
    # Set table width to full page width
    table.columns[0].width = Cm(18.0)  # Full page width minus margins
    
    return table

def generate_docx_report(form_data):
    """
    Generate a professional DOCX report that matches the client template exactly.
    Creates a form template with structured tables and blank spaces for manual entry.
    """
    doc = Document()
    
    # Set page properties
    section = doc.sections[0]
    section.page_width = Cm(21.0)  # A4 width
    section.page_height = Cm(29.7)  # A4 height
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.right_to_left = True
    
    # Section A: نموذج بطاقة الوصف المهني
    # Top title - centered, bold 20pt
    title = doc.add_heading("أ- نموذج بطاقة الوصف المهني", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Cm(0.71)  # 20pt
        run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # 1. البيانات المرجعية للمهنة
    header_table = create_header_band(doc, "1- البيانات المرجعية للمهنة")
    doc.add_paragraph()  # Spacing
    
    # Create the reference data table
    ref_table = doc.add_table(rows=7, cols=3)
    ref_table.style = 'Table Grid'
    
    # Set column widths
    set_col_widths(ref_table, [6.0, 5.0, 6.0])
    
    # Set right column labels (main labels)
    right_labels = [
        "المجموعة الرئيسية",
        "المجموعة الفرعية", 
        "المجموعة الثانوية",
        "مجموعة الوحدات",
        "المهنة",
        "موقع العمل",
        "المرتبة"
    ]
    
    # Set middle column labels (code labels)
    middle_labels = [
        "رمز المجموعة الرئيسية",
        "رمز المجموعة الفرعية",
        "رمز المجموعة الثانوية", 
        "رمز الوحدات",
        "رمز المهنة",
        "",  # blank
        ""   # blank
    ]
    
    # Fill the table
    ref_data = form_data.get('ref_data', {})
    for i in range(7):
        # Left column - values
        left_cell = ref_table.cell(i, 0)
        set_cell_borders(left_cell)
        p = left_cell.paragraphs[0]
        arabic(p)
        
        # Get value from ref_data
        value = ""
        if i == 0: value = ref_data.get('main_group', '')
        elif i == 1: value = ref_data.get('sub_group', '')
        elif i == 2: value = ref_data.get('secondary_group', '')
        elif i == 3: value = ref_data.get('unit_group', '')
        elif i == 4: value = ref_data.get('job', '')
        elif i == 5: value = ref_data.get('work_location', '')
        elif i == 6: value = ref_data.get('grade', '')
        
        p.text = value.strip() if value else ""
        
        # Middle column - code labels
        middle_cell = ref_table.cell(i, 1)
        set_cell_borders(middle_cell)
        p = middle_cell.paragraphs[0]
        arabic(p)
        p.text = middle_labels[i]
        
        # Right column - main labels
        right_cell = ref_table.cell(i, 2)
        set_cell_borders(right_cell)
        p = right_cell.paragraphs[0]
        arabic(p)
        p.text = right_labels[i]
    
    doc.add_paragraph()  # Spacing
    
    # 2. الملخص العام للمهنة
    header_table = create_header_band(doc, "2- الملخص العام للمهنة")
    doc.add_paragraph()  # Spacing
    
    # Summary table with fixed height
    summary_table = doc.add_table(rows=1, cols=1)
    summary_table.style = 'Table Grid'
    cell = summary_table.cell(0, 0)
    set_cell_borders(cell)
    
    # Set fixed height (~3.5 cm)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcHeight = OxmlElement('w:tcHeight')
    tcHeight.set(qn('w:val'), '1000')  # ~3.5 cm in twips
    tcHeight.set(qn('w:hRule'), 'exact')
    tcPr.append(tcHeight)
    
    # Add summary text if available
    summary = form_data.get('summary', '')
    if summary:
        p = cell.paragraphs[0]
        arabic(p)
        p.text = summary.strip()
    
    # Set table width
    summary_table.columns[0].width = Cm(18.0)
    
    doc.add_paragraph()  # Spacing
    
    # 3. قنوات التواصل
    header_table = create_header_band(doc, "3- قنوات التواصل")
    doc.add_paragraph()  # Spacing
    
    # Communication table
    comm_table = doc.add_table(rows=2, cols=3)
    comm_table.style = 'Table Grid'
    set_col_widths(comm_table, [6.5, 6.5, 4.0])
    
    # Row 1: Internal communications
    internal_comms = form_data.get('internal_communications', [])
    if internal_comms and len(internal_comms) > 0:
        entity = internal_comms[0].get('entity', '').strip()
        purpose = internal_comms[0].get('purpose', '').strip()
    else:
        entity = ""
        purpose = ""
    
    # Row 1
    cell1 = comm_table.cell(0, 0)  # Entity
    set_cell_borders(cell1)
    p = cell1.paragraphs[0]
    arabic(p)
    p.text = entity
    
    cell2 = comm_table.cell(0, 1)  # Purpose
    set_cell_borders(cell2)
    p = cell2.paragraphs[0]
    arabic(p)
    p.text = purpose
    
    cell3 = comm_table.cell(0, 2)  # Label
    set_cell_borders(cell3)
    p = cell3.paragraphs[0]
    arabic(p)
    p.text = "جهات التواصل الداخلية"
    
    # Row 2: External communications
    external_comms = form_data.get('external_communications', [])
    if external_comms and len(external_comms) > 0:
        entity = external_comms[0].get('entity', '').strip()
        purpose = external_comms[0].get('purpose', '').strip()
    else:
        entity = ""
        purpose = ""
    
    # Row 2
    cell1 = comm_table.cell(1, 0)  # Entity
    set_cell_borders(cell1)
    p = cell1.paragraphs[0]
    arabic(p)
    p.text = entity
    
    cell2 = comm_table.cell(1, 1)  # Purpose
    set_cell_borders(cell2)
    p = cell2.paragraphs[0]
    arabic(p)
    p.text = purpose
    
    cell3 = comm_table.cell(1, 2)  # Label
    set_cell_borders(cell3)
    p = cell3.paragraphs[0]
    arabic(p)
    p.text = "جهات التواصل الخارجية"
    
    doc.add_paragraph()  # Spacing
    
    # 4. مستويات المهنة القياسية
    header_table = create_header_band(doc, "4- مستويات المهنة القياسية")
    doc.add_paragraph()  # Spacing
    
    # Job levels table
    level_table = doc.add_table(rows=4, cols=2)
    level_table.style = 'Table Grid'
    set_col_widths(level_table, [8.5, 8.5])
    
    # Right column labels
    right_labels = [
        "مستوى المهنة القياسي",
        "رمز المستوى المهني",
        "الدور المهني",
        "التدرج المهني (المرتبة)"
    ]
    
    # Fill the table
    job_levels = form_data.get('job_levels', [])
    for i in range(4):
        # Left column - values
        left_cell = level_table.cell(i, 0)
        set_cell_borders(left_cell)
        p = left_cell.paragraphs[0]
        arabic(p)
        
        # Get value from job_levels
        value = ""
        if job_levels and len(job_levels) > 0:
            level = job_levels[0]
            if i == 0: value = level.get('level', '')
            elif i == 1: value = level.get('code', '')
            elif i == 2: value = level.get('role', '')
            elif i == 3: value = level.get('progression', '')
        
        p.text = value.strip() if value else ""
        
        # Right column - labels
        right_cell = level_table.cell(i, 1)
        set_cell_borders(right_cell)
        p = right_cell.paragraphs[0]
        arabic(p)
        p.text = right_labels[i]
    
    doc.add_paragraph()  # Spacing
    
    # 5. الجدارات
    header_table = create_header_band(doc, "5- الجدارات")
    doc.add_paragraph()  # Spacing
    
    # Competencies table
    comp_table = doc.add_table(rows=4, cols=3)
    comp_table.style = 'Table Grid'
    set_col_widths(comp_table, [9.0, 5.0, 3.0])
    
    # Fill the table
    core_comp = form_data.get('core_competencies', [])
    leadership_comp = form_data.get('leadership_competencies', [])
    technical_comp = form_data.get('technical_competencies', [])
    
    # Row 1: Basic competencies
    cell1 = comp_table.cell(0, 0)  # Value
    set_cell_borders(cell1)
    p = cell1.paragraphs[0]
    arabic(p)
    if core_comp and len(core_comp) > 0:
        p.text = core_comp[0].get('name', '').strip()
    
    cell2 = comp_table.cell(0, 1)  # Type label
    set_cell_borders(cell2)
    p = cell2.paragraphs[0]
    arabic(p)
    p.text = "الجدارات الأساسية"
    
    # Row 2: Leadership competencies
    cell1 = comp_table.cell(1, 0)  # Value
    set_cell_borders(cell1)
    p = cell1.paragraphs[0]
    arabic(p)
    if leadership_comp and len(leadership_comp) > 0:
        p.text = leadership_comp[0].get('name', '').strip()
    
    cell2 = comp_table.cell(1, 1)  # Type label
    set_cell_borders(cell2)
    p = cell2.paragraphs[0]
    arabic(p)
    p.text = "الجدارات القيادية"
    
    # Row 3: Technical competencies
    cell1 = comp_table.cell(2, 0)  # Value
    set_cell_borders(cell1)
    p = cell1.paragraphs[0]
    arabic(p)
    if technical_comp and len(technical_comp) > 0:
        p.text = technical_comp[0].get('name', '').strip()
    
    cell2 = comp_table.cell(2, 1)  # Type label
    set_cell_borders(cell2)
    p = cell2.paragraphs[0]
    arabic(p)
    p.text = "الجدارات الفنية"
    
    # Row 4: Blank
    cell1 = comp_table.cell(3, 0)  # Value
    set_cell_borders(cell1)
    cell2 = comp_table.cell(3, 1)  # Type label
    set_cell_borders(cell2)
    
    # Right column: Merge vertically and add "الجدارات السلوكية"
    merge_vertically(comp_table, 2, 0, 3)
    right_cell = comp_table.cell(0, 2)
    set_cell_borders(right_cell)
    p = right_cell.paragraphs[0]
    arabic(p)
    p.text = "الجدارات السلوكية"
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page break
    doc.add_page_break()
    
    # Section B: نموذج الوصف الفعلي
    # Top title - centered, bold 20pt
    title = doc.add_heading("ب- نموذج الوصف الفعلي", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Cm(0.71)  # 20pt
        run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # 1. المهام
    header_table = create_header_band(doc, "1- المهام")
    doc.add_paragraph()  # Spacing
    
    # Tasks table
    tasks_table = doc.add_table(rows=4, cols=1)
    tasks_table.style = 'Table Grid'
    set_col_widths(tasks_table, [18.0])
    
    # Row 1: Leadership tasks
    cell = tasks_table.cell(0, 0)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    arabic(p)
    p.text = "المهام القيادية/الإشرافية"
    
    # Add leadership tasks if available
    leadership_tasks = form_data.get('leadership_tasks', [])
    if leadership_tasks:
        for task in leadership_tasks:
            if task and task.strip():
                p = cell.add_paragraph()
                arabic(p)
                p.text = f"• {task.strip()}"
    
    # Row 2: Specialized tasks
    cell = tasks_table.cell(1, 0)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    arabic(p)
    p.text = "المهام التخصصية"
    
    # Add specialized tasks if available
    specialized_tasks = form_data.get('specialized_tasks', [])
    if specialized_tasks:
        for task in specialized_tasks:
            if task and task.strip():
                p = cell.add_paragraph()
                arabic(p)
                p.text = f"• {task.strip()}"
    
    # Row 3: Other tasks
    cell = tasks_table.cell(2, 0)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    arabic(p)
    p.text = "مهام أخرى إضافية"
    
    # Add other tasks if available
    other_tasks = form_data.get('other_tasks', [])
    if other_tasks:
        for task in other_tasks:
            if task and task.strip():
                p = cell.add_paragraph()
                arabic(p)
                p.text = f"• {task.strip()}"
    
    # Row 4: Blank spacer
    cell = tasks_table.cell(3, 0)
    set_cell_borders(cell)
    
    # Set row heights
    for row in tasks_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcHeight = OxmlElement('w:tcHeight')
            tcHeight.set(qn('w:val'), '600')  # ~1.5-2.0 cm in twips
            tcHeight.set(qn('w:hRule'), 'exact')
            tcPr.append(tcHeight)
    
    doc.add_paragraph()  # Spacing
    
    # 2. الجدارات السلوكية والفنية
    header_table = create_header_band(doc, "2- الجدارات السلوكية والفنية")
    doc.add_paragraph()  # Spacing
    
    # (a) Behavioral competencies table
    doc.add_paragraph("الجدارات السلوكية").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    behavioral_table = doc.add_table(rows=6, cols=3)  # 1 header + 5 body rows
    behavioral_table.style = 'Table Grid'
    set_col_widths(behavioral_table, [2.0, 10.0, 5.0])
    
    # Header row
    headers = ["الرقم", "الجدارات السلوكية", "مستوى الإتقان"]
    for i, header in enumerate(headers):
        cell = behavioral_table.cell(0, i)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        p.text = header
        for run in p.runs:
            run.font.bold = True
    
    # Fill data rows
    behavioral_data = form_data.get('behavioral_table', [])
    for i in range(5):
        row_idx = i + 1
        
        # Number
        cell = behavioral_table.cell(row_idx, 0)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        p.text = str(i + 1)
        
        # Competency name
        cell = behavioral_table.cell(row_idx, 1)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        if i < len(behavioral_data) and behavioral_data[i].get('name'):
            p.text = behavioral_data[i]['name'].strip()
        
        # Level
        cell = behavioral_table.cell(row_idx, 2)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        if i < len(behavioral_data) and behavioral_data[i].get('level'):
            p.text = behavioral_data[i]['level'].strip()
    
    doc.add_paragraph()  # Spacing
    
    # (b) Technical competencies table
    doc.add_paragraph("الجدارات الفنية").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    technical_table = doc.add_table(rows=6, cols=3)  # 1 header + 5 body rows
    technical_table.style = 'Table Grid'
    set_col_widths(technical_table, [2.0, 10.0, 5.0])
    
    # Header row
    headers = ["الرقم", "الجدارات الفنية", "مستوى الإتقان"]
    for i, header in enumerate(headers):
        cell = technical_table.cell(0, i)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        p.text = header
        for run in p.runs:
            run.font.bold = True
    
    # Fill data rows
    technical_data = form_data.get('technical_table', [])
    for i in range(5):
        row_idx = i + 1
        
        # Number
        cell = technical_table.cell(row_idx, 0)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        p.text = str(i + 1)
        
        # Competency name
        cell = technical_table.cell(row_idx, 1)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        if i < len(technical_data) and technical_data[i].get('name'):
            p.text = technical_data[i]['name'].strip()
        
        # Level
        cell = technical_table.cell(row_idx, 2)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        if i < len(technical_data) and technical_data[i].get('level'):
            p.text = technical_data[i]['level'].strip()
    
    doc.add_paragraph()  # Spacing
    
    # 3. إدارة الأداء المهني
    header_table = create_header_band(doc, "3- إدارة الأداء المهني")
    doc.add_paragraph()  # Spacing
    
    # KPIs table
    kpi_table = doc.add_table(rows=5, cols=3)  # 1 header + 4 body rows
    kpi_table.style = 'Table Grid'
    set_col_widths(kpi_table, [2.0, 9.0, 6.0])
    
    # Header row
    headers = ["الرقم", "مؤشرات الأداء الرئيسية", "طريقة القياس"]
    for i, header in enumerate(headers):
        cell = kpi_table.cell(0, i)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        p.text = header
        for run in p.runs:
            run.font.bold = True
    
    # Fill data rows
    kpis = form_data.get('kpis', [])
    for i in range(4):
        row_idx = i + 1
        
        # Number
        cell = kpi_table.cell(row_idx, 0)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        p.text = str(i + 1)
        
        # KPI metric
        cell = kpi_table.cell(row_idx, 1)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        if i < len(kpis) and kpis[i].get('metric'):
            p.text = kpis[i]['metric'].strip()
        
        # Measurement method
        cell = kpi_table.cell(row_idx, 2)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        arabic(p)
        if i < len(kpis) and kpis[i].get('measure'):
            p.text = kpis[i]['measure'].strip()
    
    return doc
